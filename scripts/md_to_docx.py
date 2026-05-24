#!/usr/bin/env python3
"""Convert docs/gameplan.md to docs/gameplan.docx with basic markdown support."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run_elem.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run_elem.append(text_elem)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\)|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if m:
                add_hyperlink(paragraph, m.group(2), m.group(1))
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.LINE)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for col, header in enumerate(headers):
                cell = table.rows[0].cells[col]
                cell.text = header
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row):
                    if c_idx < len(headers):
                        table.rows[r_idx + 1].cells[c_idx].text = value
            doc.add_paragraph()
            continue

        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines):
                item = lines[i].strip()
                if not re.match(r"^[-*]\s+", item):
                    if item.startswith("  ") and lines[i - 1].strip().startswith("- ["):
                        pass
                    else:
                        break
                if re.match(r"^[-*]\s+\[[ xX]\]", item):
                    checked = "[x]" in item.lower()[:6]
                    text = re.sub(r"^[-*]\s+\[[ xX]\]\s*", "", item)
                    p = doc.add_paragraph(style="List Bullet")
                    prefix = "☑ " if checked else "☐ "
                    add_rich_text(p, prefix + text)
                elif re.match(r"^[-*]\s+", item):
                    text = re.sub(r"^[-*]\s+", "", item)
                    p = doc.add_paragraph(style="List Bullet")
                    add_rich_text(p, text)
                i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                text = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_rich_text(p, text)
                i += 1
            continue

        p = doc.add_paragraph()
        add_rich_text(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "gameplan.md"
    out = root / "docs" / "gameplan.docx"
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out = Path(sys.argv[2])
    convert(md, out)
